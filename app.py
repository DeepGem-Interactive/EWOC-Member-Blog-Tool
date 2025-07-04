from flask import Flask, render_template, request, redirect, url_for, session, send_from_directory, jsonify, g, send_file
from docx import Document
from openai import AzureOpenAI
from dotenv import load_dotenv
import os
import time
from datetime import datetime, timedelta
import requests
import markdown
from bs4 import BeautifulSoup
from flask_session import Session
import json
import pyodbc
from io import BytesIO
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import hashlib
import secrets
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
import sqlite3
import base64
import io
from PIL import Image
import uuid
from urllib.parse import unquote

load_dotenv()

app = Flask(__name__, static_folder='static')
app.secret_key = os.getenv('FLASK_SECRET_KEY')
app.config['UPLOAD_FOLDER'] = 'static/generated'

# Database configuration
app.config['AZURE_SQL_SERVER'] = os.getenv('AZURE_SQL_SERVER')
app.config['AZURE_SQL_DATABASE'] = os.getenv('AZURE_SQL_DATABASE')
app.config['AZURE_SQL_USERNAME'] = os.getenv('AZURE_SQL_USERNAME')
app.config['AZURE_SQL_PASSWORD'] = os.getenv('AZURE_SQL_PASSWORD')

# Session configuration
app.config['SESSION_TYPE'] = 'filesystem'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=1)
app.config['SESSION_COOKIE_SECURE'] = True
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
Session(app)

def get_db():
    if 'db' not in g:
        conn_str = (
            f"DRIVER={{ODBC Driver 18 for SQL Server}};"
            f"SERVER={app.config['AZURE_SQL_SERVER']};"
            f"DATABASE={app.config['AZURE_SQL_DATABASE']};"
            f"UID={app.config['AZURE_SQL_USERNAME']};"
            f"PWD={app.config['AZURE_SQL_PASSWORD']};"
            "Encrypt=yes;TrustServerCertificate=no;"
        )
        g.db = pyodbc.connect(conn_str)
    return g.db

def close_db():
    db = g.pop('db', None)
    if db is not None:
        db.close()

def init_db():
    with app.app_context():
        db = get_db()
        cursor = db.cursor()
        
        # Create users table if it doesn't exist
        cursor.execute('''
        IF NOT EXISTS (SELECT * FROM sys.tables WHERE name = 'users')
        CREATE TABLE users (
            id INT IDENTITY(1,1) PRIMARY KEY,
            username NVARCHAR(255) UNIQUE NOT NULL,
            email NVARCHAR(255) UNIQUE NOT NULL,
            password NVARCHAR(500) NOT NULL,
            firm NVARCHAR(255),
            location NVARCHAR(255),
            lawyer_name NVARCHAR(255),
            state NVARCHAR(50),
            address NVARCHAR(255),
            planning_session NVARCHAR(255),
            other_planning_session NVARCHAR(255),
            discovery_call_link NVARCHAR(255)
        )
        ''')
        
        # Add new columns if they don't exist
        try:
            cursor.execute('''
            IF NOT EXISTS (SELECT * FROM sys.columns WHERE object_id = OBJECT_ID('users') AND name = 'address')
            ALTER TABLE users ADD address NVARCHAR(255)
            ''')
        except pyodbc.Error:
            pass
            
        try:
            cursor.execute('''
            IF NOT EXISTS (SELECT * FROM sys.columns WHERE object_id = OBJECT_ID('users') AND name = 'planning_session')
            ALTER TABLE users ADD planning_session NVARCHAR(255)
            ''')
        except pyodbc.Error:
            pass
            
        try:
            cursor.execute('''
            IF NOT EXISTS (SELECT * FROM sys.columns WHERE object_id = OBJECT_ID('users') AND name = 'other_planning_session')
            ALTER TABLE users ADD other_planning_session NVARCHAR(255)
            ''')
        except pyodbc.Error:
            pass
            
        try:
            cursor.execute('''
            IF NOT EXISTS (SELECT * FROM sys.columns WHERE object_id = OBJECT_ID('users') AND name = 'discovery_call_link')
            ALTER TABLE users ADD discovery_call_link NVARCHAR(255)
            ''')
        except pyodbc.Error:
            pass
        
        # Update password column size to accommodate hashed passwords
        try:
            cursor.execute('''
            ALTER TABLE users ALTER COLUMN password NVARCHAR(500)
            ''')
        except pyodbc.Error:
            pass
        
        # Create tones table if it doesn't exist
        cursor.execute('''
        IF NOT EXISTS (SELECT * FROM sys.tables WHERE name = 'tones')
        CREATE TABLE tones (
            id INT IDENTITY(1,1) PRIMARY KEY,
            user_id INT NOT NULL,
            name NVARCHAR(255) NOT NULL,
            description NVARCHAR(MAX) NOT NULL,
            CONSTRAINT UQ_user_tone UNIQUE(user_id, name),
            CONSTRAINT FK_user_tone FOREIGN KEY(user_id) REFERENCES users(id)
        )
        ''')
        
        # Check if default users exist
        cursor.execute("SELECT * FROM users WHERE username IN ('admin', 'memberhub')")
        existing_users = cursor.fetchall()
        existing_usernames = [user.username for user in existing_users]
        
        if 'admin' not in existing_usernames:
            cursor.execute('''
            INSERT INTO users (username, email, password, firm, location, lawyer_name, state, address, 
                             planning_session, other_planning_session, discovery_call_link)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                'admin', 
                'admin@lawfirm.com', 
                'password123', 
                'Legal Partners', 
                'New York', 
                'John', 
                'NY',
                '',
                '',
                '',
                ''
            ))
        
        if 'memberhub' not in existing_usernames:
            cursor.execute('''
            INSERT INTO users (username, email, password, firm, location, lawyer_name, state, address, 
                             planning_session, other_planning_session, discovery_call_link)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                'memberhub', 
                'memberhub@newlawbusinessmodel.com', 
                'memberhub123', 
                'New Law Business Model', 
                'Global', 
                'Member Hub', 
                'CA',
                '',
                '',
                '',
                ''
            ))
        
        # Create password_resets table if it doesn't exist
        cursor.execute('''
        IF NOT EXISTS (SELECT * FROM sysobjects WHERE name='password_resets' AND xtype='U')
        CREATE TABLE password_resets (
            id INTEGER IDENTITY(1,1) PRIMARY KEY,
            email NVARCHAR(255) NOT NULL,
            token NVARCHAR(255) NOT NULL UNIQUE,
            expires DATETIME NOT NULL,
            used INTEGER DEFAULT 0,
            created_at DATETIME DEFAULT GETDATE()
        )
        ''')
        
        # Create feedback table if it doesn't exist
        cursor.execute('''
        IF NOT EXISTS (SELECT * FROM sys.tables WHERE name = 'feedback')
        CREATE TABLE feedback (
            id INT IDENTITY(1,1) PRIMARY KEY,
            user_id INT,
            feedback_type NVARCHAR(50) NOT NULL,
            priority NVARCHAR(20) NOT NULL,
            subject NVARCHAR(255) NOT NULL,
            message NVARCHAR(MAX) NOT NULL,
            contact_email NVARCHAR(255),
            status NVARCHAR(20) DEFAULT 'pending',
            created_at DATETIME DEFAULT GETDATE(),
            updated_at DATETIME DEFAULT GETDATE(),
            CONSTRAINT FK_feedback_user FOREIGN KEY(user_id) REFERENCES users(id)
        )
        ''')
        
        db.commit()

# Initialize database
with app.app_context():
    init_db()

# Add context processor to inject current year into all templates
@app.context_processor
def inject_year():
    return {'now': datetime.now()}

class UserSession:
    @staticmethod
    def register(email, password, firm, location, lawyer_name, state, address="", planning_session="", other_planning_session="", discovery_call_link=""):
        db = get_db()
        username = email.split('@')[0].lower()
        try:
            cursor = db.cursor()
            cursor.execute('''
            INSERT INTO users (username, email, password, firm, location, lawyer_name, state, address, 
                             planning_session, other_planning_session, discovery_call_link)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (username, email, password, firm, location, lawyer_name, state, address, 
                 planning_session, other_planning_session, discovery_call_link))
            db.commit()
            return True
        except pyodbc.IntegrityError:
            return False

    @staticmethod
    def login(email, password):
        db = get_db()
        username = email.split('@')[0].lower()
        cursor = db.cursor()
        cursor.execute('SELECT * FROM users WHERE username = ?', (username,))
        user = cursor.fetchone()
        
        if user and user.password == password:
            # Get user's custom tones
            cursor.execute('SELECT name, description FROM tones WHERE user_id = ?', (user.id,))
            tones = cursor.fetchall()
            
            session['user'] = {
                'id': user.id,
                'username': user.username,
                'email': user.email,
                'firm': user.firm,
                'location': user.location,
                'lawyer_name': user.lawyer_name,
                'state': user.state,
                'address': user.address,
                'planning_session': user.planning_session,
                'other_planning_session': user.other_planning_session,
                'discovery_call_link': user.discovery_call_link,
                'custom_tones': [{'name': tone.name, 'description': tone.description} for tone in tones]
            }
            return True
        return False

    @staticmethod
    def update_profile(username, firm, location, lawyer_name, state, address="", planning_session="", other_planning_session="", discovery_call_link=""):
        db = get_db()
        try:
            cursor = db.cursor()
            cursor.execute('''
            UPDATE users 
            SET firm = ?, location = ?, lawyer_name = ?, state = ?, 
                address = ?, planning_session = ?, other_planning_session = ?, discovery_call_link = ?
            WHERE username = ?
            ''', (firm, location, lawyer_name, state, address, planning_session, 
                 other_planning_session, discovery_call_link, username))
            db.commit()
            
            # Update session if this is the current user
            if 'user' in session and session['user']['username'] == username:
                session['user'].update({
                    'firm': firm,
                    'location': location,
                    'lawyer_name': lawyer_name,
                    'state': state,
                    'address': address,
                    'planning_session': planning_session,
                    'other_planning_session': other_planning_session,
                    'discovery_call_link': discovery_call_link
                })
                session.modified = True
            return True
        except pyodbc.Error:
            return False

    @staticmethod
    def get_current_user():
        return session.get('user')

    @staticmethod
    def add_custom_tone(user_id, tone_name, tone_description):
        db = get_db()
        try:
            cursor = db.cursor()
            cursor.execute('''
            INSERT INTO tones (user_id, name, description)
            VALUES (?, ?, ?)
            ''', (user_id, tone_name, tone_description))
            db.commit()
            
            # Update session if this is the current user
            if 'user' in session and session['user']['id'] == user_id:
                session['user']['custom_tones'].append({
                    'name': tone_name,
                    'description': tone_description
                })
                session.modified = True
            return True
        except pyodbc.IntegrityError:
            return False
    
    @staticmethod
    def get_custom_tones(user_id):
        db = get_db()
        cursor = db.cursor()
        cursor.execute('SELECT name, description FROM tones WHERE user_id = ?', (user_id,))
        tones = cursor.fetchall()
        return [{'name': tone.name, 'description': tone.description} for tone in tones]

    @staticmethod
    def submit_feedback(user_id, feedback_type, priority, subject, message, contact_email=None):
        db = get_db()
        try:
            cursor = db.cursor()
            cursor.execute('''
            INSERT INTO feedback (user_id, feedback_type, priority, subject, message, contact_email)
            VALUES (?, ?, ?, ?, ?, ?)
            ''', (user_id, feedback_type, priority, subject, message, contact_email))
            db.commit()
            return True
        except Exception as e:
            print(f"Error submitting feedback: {str(e)}")
            return False

class Config:
    ARTICLES_DIR = "articles"
    GENERATED_DIR = "generated"
    os.makedirs(ARTICLES_DIR, exist_ok=True)
    os.makedirs(GENERATED_DIR, exist_ok=True)

    # Default section markers (can be updated based on client requirements)
    SECTION_MARKERS = {
        'hook': {
            'start': 0,  # First paragraph
            'end': 1     # End of first paragraph
        },
        'summary': {
            'start': 1,  # Second paragraph (2-3 lines ending with "read more...")
            'end': 2     # End of second paragraph
        },
        'disclaimer': {
            'start': -1,  # Last paragraph (disclaimer)
            'end': None   # End of content
        }
    }

class AzureServices:
    def __init__(self):
        self.text_client = AzureOpenAI(
            api_key=os.getenv("AZURE_OPENAI_KEY"),
            api_version="2024-02-15-preview",
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
        )
        
        self.conversations = {}

    def _extract_sections(self, content):
        """Extract and preserve specific sections from the content."""
        try:
            paragraphs = content.split('\n\n')
            preserved_sections = {}
            
            print(f"\n=== EXTRACTING SECTIONS ===")
            print(f"Total paragraphs found: {len(paragraphs)}")
            
            # Extract Hook (first paragraph)
            hook_start = Config.SECTION_MARKERS['hook']['start']
            hook_end = Config.SECTION_MARKERS['hook']['end']
            if hook_start < len(paragraphs):
                preserved_sections['hook'] = paragraphs[hook_start:hook_end][0] if hook_end - hook_start == 1 else '\n\n'.join(paragraphs[hook_start:hook_end])
                print(f"✓ Extracted Hook (paragraph 1): {preserved_sections['hook'][:100]}...")
            else:
                print("✗ Warning: Hook section not found in content")
                preserved_sections['hook'] = ""
            
            # Extract Summary (second paragraph - should be 2-3 lines ending with "read more...")
            summary_start = Config.SECTION_MARKERS['summary']['start']
            summary_end = Config.SECTION_MARKERS['summary']['end']
            if summary_start < len(paragraphs):
                preserved_sections['summary'] = paragraphs[summary_start:summary_end][0] if summary_end - summary_start == 1 else '\n\n'.join(paragraphs[summary_start:summary_end])
                print(f"✓ Extracted Summary (paragraph 2): {preserved_sections['summary'][:100]}...")
            else:
                print("✗ Warning: Summary section not found in content")
                preserved_sections['summary'] = ""
            

            
            # Extract Disclaimer (last paragraph)
            disclaimer_start = Config.SECTION_MARKERS['disclaimer']['start']
            if len(paragraphs) > 0:
                preserved_sections['disclaimer'] = paragraphs[disclaimer_start]
                print(f"✓ Extracted Disclaimer (paragraph {len(paragraphs)}): {preserved_sections['disclaimer'][:100]}...")
            else:
                print("✗ Warning: Disclaimer section not found in content")
                preserved_sections['disclaimer'] = ""
            
            print("=== SECTION EXTRACTION COMPLETE ===\n")
            
            return preserved_sections
        except Exception as e:
            print(f"Error extracting sections: {str(e)}")
            return {'hook': "", 'summary': "", 'disclaimer': ""}

    def _reconstruct_content(self, new_content, preserved_sections):
        """Preserve specific sections exactly as they are in the original content."""
        try:
            paragraphs = new_content.split('\n\n')
            
            print(f"\n=== RECONSTRUCTING CONTENT ===")
            print(f"Total paragraphs in new content: {len(paragraphs)}")
            
            # Remove any existing disclaimer or service description from the middle of the content
            # to prevent duplication, but preserve CTA content
            cleaned_paragraphs = []
            for i, para in enumerate(paragraphs):
                # Skip paragraphs that are exact matches of disclaimers or service descriptions
                # But preserve paragraphs that contain CTA phrases
                is_duplicate_section = False
                if preserved_sections['disclaimer'] and preserved_sections['disclaimer'].strip() == para.strip():
                    is_duplicate_section = True
                    print(f"⚠️  Removed duplicate disclaimer from position {i+1}")

                
                if not is_duplicate_section:
                    cleaned_paragraphs.append(para)
            
            paragraphs = cleaned_paragraphs
            print(f"Paragraphs after cleaning: {len(paragraphs)}")
            
            # Place preserved sections in their correct positions
            # Hook (first paragraph)
            if preserved_sections['hook'] and len(paragraphs) > 0:
                paragraphs[0] = preserved_sections['hook']
                print("✓ Preserved Hook section at position 1")
            
            # Summary (second paragraph - 2-3 lines ending with "read more...")
            if preserved_sections['summary'] and len(paragraphs) > 1:
                paragraphs[1] = preserved_sections['summary']
                print("✓ Preserved Summary section at position 2")
            elif preserved_sections['summary'] and len(paragraphs) <= 1:
                # If there aren't enough paragraphs, add the summary as second paragraph
                if len(paragraphs) == 0:
                    paragraphs.append("")  # Add empty first paragraph if needed
                paragraphs.append(preserved_sections['summary'])
                print("✓ Added Summary section at position 2")
            

            
            # Disclaimer (last paragraph)
            if preserved_sections['disclaimer']:
                # Always add disclaimer at the very end
                paragraphs.append(preserved_sections['disclaimer'])
                print("✓ Added Disclaimer section at the very end")
            
            final_content = '\n\n'.join(paragraphs)
            
            # Verify sections are preserved exactly
            if preserved_sections['hook'] and preserved_sections['hook'] not in final_content:
                print("✗ Warning: Hook section not preserved exactly")
            if preserved_sections['summary'] and preserved_sections['summary'] not in final_content:
                print("✗ Warning: Summary section not preserved exactly")

            if preserved_sections['disclaimer'] and preserved_sections['disclaimer'] not in final_content:
                print("✗ Warning: Disclaimer section not preserved exactly")
            
            print("=== RECONSTRUCTION COMPLETE ===\n")
            
            return final_content
        except Exception as e:
            print(f"Error preserving sections: {str(e)}")
            return new_content

    def _validate_with_gpt(self, original_text, new_content, components):
        """Validate article components using GPT for better semantic understanding."""
        validation_prompt = f"""
            You are an expert content validator. Analyze these two articles and provide a detailed validation.
            You MUST respond with a valid JSON object following this EXACT structure, with no additional text:
            
            {{
                "components": {{
                    "keywords": {{
                        "found": true/false,
                        "occurrences": number,
                        "variations": ["variation1", "variation2"],
                        "in_first_150": true/false
                    }},
                    "firm_info": {{
                        "found": true/false,
                        "name": true/false,
                        "location": true/false
                    }},
                    "lawyer_info": {{
                        "found": true/false,
                        "name": true/false,
                        "location": true/false
                    }},
                    "planning_session": {{
                        "found": true/false,
                        "name": true/false,
                        "references": number
                    }},
                    "discovery_call": {{
                        "found": true/false,
                        "link": true/false,
                        "references": number
                    }}
                }},
                "preserved_sections": {{
                    "hook": true/false,
                    "summary": true/false,
                    "disclaimer": true/false
                }},
                "change_analysis": {{
                    "percentage": number,
                    "significant_changes": true/false,
                    "maintained_essence": true/false
                }},
                "warnings": ["warning1", "warning2"],
                "missing_components": ["component1", "component2"]
            }}

            Analyze the following content:

            Required components to check:
            - Keywords: {components['keywords']}
            - Firm: {components['firm_name']} in {components['location']}
            - Lawyer: {components['lawyer_name']} in {components['city']}, {components['state']}
            - Planning Session: {components['planning_session_name']}
            - Discovery Call: {components['discovery_call_link']}

            Original Article:
            {original_text}

            New Article:
            {new_content}

            Remember to respond with ONLY the JSON object, no additional text or explanation.
        """

        try:
            response = self.text_client.chat.completions.create(
                model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
                messages=[
                    {"role": "system", "content": "You are a JSON-only response validator. Always respond with valid JSON matching the exact structure provided."},
                    {"role": "user", "content": validation_prompt}
                ],
                temperature=0.1,  # Lower temperature for more consistent JSON output
                response_format={ "type": "json_object" }  # Force JSON response
            )

            # Get the response content and ensure it's valid JSON
            response_content = response.choices[0].message.content.strip()
            
            # Try to parse the JSON response
            try:
                validation_results = json.loads(response_content)
            except json.JSONDecodeError as e:
                print(f"Error parsing JSON response: {str(e)}")
                print(f"Raw response: {response_content}")
                raise

            # Validate the structure of the response
            required_keys = ['components', 'preserved_sections', 'change_analysis', 'warnings', 'missing_components']
            if not all(key in validation_results for key in required_keys):
                print("Invalid response structure. Missing required keys.")
                raise ValueError("Invalid response structure")

            # Print validation results in a readable format
            print("\n=== GPT Article Validation Results ===")
            
            print("\nComponent Status:")
            for component, details in validation_results['components'].items():
                status = '✓' if details.get('found', False) else '✗'
                print(f"- {component}: {status}")
                if component == 'keywords' and details.get('variations'):
                    print(f"  • Variations found: {', '.join(details['variations'])}")
                if 'occurrences' in details:
                    print(f"  • Occurrences: {details['occurrences']}")

            print("\nPreserved Sections:")
            for section, preserved in validation_results['preserved_sections'].items():
                print(f"- {section}: {'✓' if preserved else '✗'}")

            print(f"\nChange Analysis:")
            print(f"- Change Percentage: {validation_results['change_analysis']['percentage']:.1f}%")
            print(f"- Significant Changes: {'✓' if validation_results['change_analysis']['significant_changes'] else '✗'}")
            print(f"- Maintained Essence: {'✓' if validation_results['change_analysis']['maintained_essence'] else '✗'}")

            if validation_results['warnings']:
                print("\nWarnings:")
                for warning in validation_results['warnings']:
                    print(f"- {warning}")

            if validation_results['missing_components']:
                print("\nMissing Components:")
                for component in validation_results['missing_components']:
                    print(f"- {component}")

            print("===============================\n")

            return validation_results

        except Exception as e:
            print(f"Error in GPT validation: {str(e)}")
            print("Unable to validate article components. Please check the generated content manually.")
            return None

    def rewrite_content(self, original_text, tone, tone_description, keywords, firm_name, location, lawyer_name, city, state, discovery_call_link, planning_session_name="Life & Legacy Planning Session"):
        try:
            # Extract sections to preserve
            print("\nExtracting sections to preserve...")
            preserved_sections = self._extract_sections(original_text)
            
            # CRITICAL: DO NOT MODIFY THESE SECTIONS {preserved_sections}:
            # 1. The first paragraph (Hook) - which is this
            # 2. The fourth paragraph (Plug) - Keep it exactly as is
            # 3. The last paragraph (Disclaimer) - Keep it exactly as is
            
            # CRITICAL: DO NOT REPEAT THESE SECTIONS:
            # 1. The hook paragraph should appear ONLY ONCE at the beginning
            # 2. The plug paragraph should appear ONLY ONCE in its original position
            # 3. The disclaimer paragraph should appear ONLY ONCE at the end
            # 4. DO NOT include these preserved sections anywhere else in the article
            # 5. DO NOT create new paragraphs that repeat the same content as the hook, plug, or disclaimer

            # Add explicit instructions about preserving sections
            # CRITICAL: DO NOT REPEAT THESE SECTIONS:
            # 1. The first paragraph should appear ONLY ONCE at the beginning
            # 2. The second paragraph should appear ONLY ONCE in its original position
            # 3. The disclaimer paragraph should appear ONLY ONCE at the end
            # 4. DO NOT include these preserved sections anywhere else in the article
            # 5. DO NOT create new paragraphs that repeat the same content as the hook, plug, or disclaimer


            system_prompt = f"""
                You are a legal blog post rewriter. There should be At least 40% changes from original. Rewrite the article following these strict guidelines:
                
                CRITICAL: DO NOT MODIFY THESE SECTIONS {preserved_sections}:
                1. The first paragraph (Hook): {preserved_sections['hook']}
                   - This should remain exactly as is and NOT be duplicated anywhere else
                
                2. The second paragraph (Summary): {preserved_sections['summary']}
                   - This should remain exactly as is and NOT be duplicated anywhere else
                   - This should be 2-3 lines ending with "read more..."
                
                3. The last paragraph (Disclaimer): {preserved_sections['disclaimer']}
                   - This should remain exactly as is and NOT be duplicated anywhere else
                   - This is the legal disclaimer paragraph
                
                EXPECTED ARTICLE STRUCTURE:
                1. Hook paragraph (preserved)
                2. Summary paragraph (preserved - 2-3 lines ending with "read more...")
                3. Main heading/title (starts with #)
                4. Article content (rewrite this part)
                5. CTA (call-to-action)
                6. Disclaimer paragraph (preserved)
                
                CRITICAL GENERATION ORDER:
                - Generate content in this EXACT order: Hook → Summary → Heading → Article Content → CTA → Disclaimer
                - The main heading/title MUST come immediately after the summary paragraph
                - Start your generated content with the heading (e.g., "# Understanding Pet Trusts in Texas")
                - Include the CTA at the end of your content, before the disclaimer
                - The AI should generate the content in the correct order from the beginning
                
                CRITICAL: DO NOT include the disclaimer paragraph anywhere in the middle of the article.
                The disclaimer should ONLY appear at the very end of the article.
                Generate the main article content (section 3-5) in the correct order - the preserved sections will be added automatically.
                
                SEO REQUIREMENTS:
                1. Must include these elements:
                   - Primary keywords: {keywords}
                   - Firm name: {firm_name}
                   - City-state of firm: {location}
                   - Lawyer name: {lawyer_name}
                   - City-state of Lawyer: {city}, {state}
                   - Planning session name: {planning_session_name}
                   - Discovery call link: {discovery_call_link}
                2. Incorporate naturally - don't just list them
                
                TONE REQUIREMENTS:
                1. Primary Tone: {tone}
                2. Tone Description: {tone_description}
                3. Consistency: Maintain this tone throughout the entire article
                
                SPECIAL BRANDING REQUIREMENTS:
                - Avoid transactional language like "investing in" which are not aligned with the Personal Family Lawyer brand tone
                - Instead use phrases like:
                    * "work with us to choose a plan that works to keep your loved ones out of court and out of conflict"
                    * "create a plan that protects what matters most"
                    * "develop a comprehensive approach to safeguarding your family's future"
                    * "put a plan in place that ensures your wishes are honored"
                    * "create a plan that grows with your family and ensures lasting peace of mind"
                - Emphasize the ongoing relationship and family protection aspects rather than transactional terms
                - Use the term "{planning_session_name}" when referencing to planning sessions.

                CONTENT GUIDELINES:
                DO's:
                1. Use active voice
                2. Structure with 5 sections: introduction, 3 subheadings
                3. Keep length between 1000-1200 words and the summary (second paragraph) should'nt be more than 2-3 lines
                4. Use transition sentences between sections
                5. Include 1-2 bulleted lists in the entire article
                6. Balance paragraphs and lists appropriately
                7. Write in a {tone} tone
                8. Include these keywords naturally: {keywords}
                9. Mention {firm_name} in {location} where relevant
                10. Firm name is {firm_name} and location is {location}
                11 Lawyer name is {lawyer_name} and location is {city}, {state}
                12. The meeting scheduling link is {discovery_call_link}
                13. The planning session name is {planning_session_name}
                14. Make sure to include the following dynamic components where ever required:
                     - keywords: {keywords}
                     - firm_name: {firm_name}
                     - location: {location}
                     - lawyer_name: {lawyer_name}
                     - city: {city}
                     - state: {state}
                     - planning_session_name: {planning_session_name}
                     - discovery_call_link: {discovery_call_link}
                15. Generate content in this EXACT order: Heading → Article Content → CTA
                
                DON'Ts:
                1. Avoid legal jargon or complex language (keep it high-school level)
                2. No passive voice
                3. Don't use lists without context
                4. Limit metaphors
                5. Don't make conclusion too long
                6. Don't include more than 5 sources
                7. Don't exceed 1200 words
                8. Don't use more than 3 lists
                9. Don't repeat any paragraph meaning no same paragraph should be present    
                10. Don't extend the summary which is the second paragraph should'nt be more than 2-3 lines

                CTA REQUIREMENTS:
                1. MUST use the exact phrase "15-minute Discovery Call" (never "consultation" or "consult")
                2. Standard format: "Schedule your complimentary 15-minute Discovery Call with {firm_name} today"
                3. Include a clear call-to-action like "Click here to schedule" or "Book your Discovery Call now"
                4. Never offer to answer questions or provide consultation during this call
                5. CRITICAL: After the call-to-action, DO NOT add any additional paragraphs or content
                6. The article should end immediately after the call-to-action - no extra content
                
                Formatting Requirements:
                # Main Title
                ## Subheading 1
                ### Sub-subheading (if needed)
                **Bold important terms**
                - Bullet points when appropriate
                [Link text](URL) for references
                
                The article must be valuable, engaging, and optimized for both readers and search engines.
            """
            
            print("\nGenerating rewritten content...")
            response = self.text_client.chat.completions.create(
                model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": original_text}
                ],
                temperature=0.7,
            )
            
            # Get the rewritten content
            rewritten_content = response.choices[0].message.content
            
            # Preserve the sections exactly as they are
            print("\nPreserving sections exactly as they are...")
            final_content = self._reconstruct_content(rewritten_content, preserved_sections)
            
            # Validate and cleanup the structure to ensure exact format
            print("\nValidating and cleaning article structure...")
            final_content = self._validate_and_cleanup_structure(final_content, preserved_sections)
            
            # Validate the generated content using GPT
            components = {
                'keywords': keywords,
                'firm_name': firm_name,
                'location': location,
                'lawyer_name': lawyer_name,
                'city': city,
                'state': state,
                'planning_session_name': planning_session_name,
                'discovery_call_link': discovery_call_link
            }
            
            validation_results = self._validate_with_gpt(original_text, final_content, components)
            
            if validation_results is None:
                print("Warning: Article validation failed. Please review the content manually.")
            
            print("\nArticle generation complete!")
            return final_content
            
        except Exception as e:
            print(f"Error in rewrite_content: {str(e)}")
            return original_text

    def edit_content(self, session_id, user_message, current_content=None):
        if session_id not in self.conversations:
            self.conversations[session_id] = [
                {"role": "system", "content": """
                    You are a legal blog post editor. When the user requests changes:
                    1. The first paragraph should not be repeating and be the same as the original
                    2. The second paragraph should not be extending and be the same as the original
                    3. The last paragraph (disclaimer) should not be repeating and be the same as the original
                    1. Make ONLY the requested changes
                    2. Return the COMPLETE updated blog (not just updated part) in markdown format
                    3. Don't include any commentary or explanations
                    4. Preserve all formatting and structure
                    5. Don't repeat any paragraph meaning no same paragraph should be present.
                """}
            ]
        
        if current_content:
            self.conversations[session_id].append(
                {"role": "assistant", "content": current_content}
            )
        
        self.conversations[session_id].append(
            {"role": "user", "content": user_message}
        )
        
        response = self.text_client.chat.completions.create(
            model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
            messages=self.conversations[session_id],
            temperature=0.5
        )
        
        ai_response = response.choices[0].message.content
        self.conversations[session_id].append(
            {"role": "assistant", "content": ai_response}
        )
        
        return ai_response
    
    def _validate_and_cleanup_structure(self, content, preserved_sections):
        """Validate and cleanup the article structure to ensure it follows the exact format."""
        try:
            print(f"\n=== VALIDATING AND CLEANING ARTICLE STRUCTURE ===")
            
            paragraphs = content.split('\n\n')
            cleaned_paragraphs = []
            
            # Step 1: Ensure hook is first
            if preserved_sections['hook'] and len(paragraphs) > 0:
                cleaned_paragraphs.append(preserved_sections['hook'])
                print("✓ Hook placed at position 1")
            else:
                print("⚠️  No hook found")
            
            # Step 2: Add line break after hook
            cleaned_paragraphs.append("")
            print("✓ Added line break after hook")
            
            # Step 3: Ensure summary is second (2-3 lines)
            if preserved_sections['summary']:
                cleaned_paragraphs.append(preserved_sections['summary'])
                print("✓ Summary placed at position 3")
            else:
                print("⚠️  No summary found")
            
            # Step 4: Find and organize content (heading, article content, CTA)
            cta_found = False
            heading_found = False
            content_paragraphs = []
            
            for i, para in enumerate(paragraphs):
                # Skip if this is the hook or summary (already handled)
                if para.strip() == preserved_sections.get('hook', '').strip():
                    continue
                if para.strip() == preserved_sections.get('summary', '').strip():
                    continue
                
                # Check if this paragraph contains CTA
                cta_phrases = ["Click here to schedule", "Book your Discovery Call", "Schedule your complimentary"]
                if any(phrase in para for phrase in cta_phrases):
                    cta_found = True
                    content_paragraphs.append(para)  # Include the CTA paragraph
                    print(f"✓ Found CTA at paragraph {i+1}")
                    break
                
                # Check if this is a heading (starts with #)
                if para.strip().startswith('#'):
                    if not heading_found:
                        heading_found = True
                        print(f"✓ Found heading at paragraph {i+1}: {para.strip()[:50]}...")
                    content_paragraphs.append(para)
                else:
                    # Add all other content paragraphs
                    content_paragraphs.append(para)
            
            # Add all content paragraphs (including heading and main content)
            cleaned_paragraphs.extend(content_paragraphs)
            
            if not heading_found:
                print("⚠️  No heading found in content")
            if not cta_found:
                print("⚠️  No CTA found in content")
            
            # Step 5: Add disclaimer at the end (only if it exists)
            if preserved_sections['disclaimer']:
                cleaned_paragraphs.append(preserved_sections['disclaimer'])
                print("✓ Disclaimer placed at the end")
            else:
                print("⚠️  No disclaimer found")
            
            # Step 6: Add any remaining content after disclaimer (preserve all content)
            for para in paragraphs:
                # Skip if this is the hook, summary, or disclaimer (already handled)
                if para.strip() == preserved_sections.get('hook', '').strip():
                    continue
                if para.strip() == preserved_sections.get('summary', '').strip():
                    continue
                if para.strip() == preserved_sections.get('disclaimer', '').strip():
                    continue
                
                # Check if this paragraph is already in content_paragraphs
                if para not in content_paragraphs:
                    cleaned_paragraphs.append(para)
                    print(f"✓ Added remaining content: {para.strip()[:50]}...")
            
            final_content = '\n\n'.join(cleaned_paragraphs)
            
            print(f"=== STRUCTURE VALIDATION COMPLETE ===")
            print(f"Final structure: Hook → Line Break → Summary → Heading → Content → CTA → Disclaimer → All Remaining Content")
            print(f"Total paragraphs: {len(cleaned_paragraphs)}")
            
            return final_content
            
        except Exception as e:
            print(f"Error in structure validation: {str(e)}")
            return content

class ImageGenerator:
    def __init__(self):
        self.image_client = AzureOpenAI(
            api_key=os.getenv("AZURE_DALLE_KEY"),
            api_version="2024-02-01",
            azure_endpoint=os.getenv("AZURE_DALLE_ENDPOINT")
        )
        self.text_client = AzureOpenAI(
            api_key=os.getenv("AZURE_OPENAI_KEY"),
            api_version="2024-02-15-preview",
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT")
        )

    def generate_image(self, text_prompt):
        try:
            safe_prompt = self._get_safe_image_prompt(text_prompt)
            
            response = self.image_client.images.generate(
                model=os.getenv("AZURE_DALLE_DEPLOYMENT"),
                prompt=safe_prompt,
                size="1024x1024",
                quality="standard",
                n=1,
            )
            image_url = response.data[0].url
            os.makedirs(os.path.join(app.static_folder, 'generated'), exist_ok=True)
            
            timestamp = int(time.time())
            image_filename = f"image_{timestamp}.png"
            image_path = os.path.join(app.static_folder, 'generated', image_filename)
            
            response = requests.get(image_url)
            with open(image_path, 'wb') as f:
                f.write(response.content)
            
            return image_filename
            
        except Exception as e:
            print(f"Image generation failed: {e}")
            return None
        
    def _get_safe_image_prompt(self, text_prompt):
        response = self.text_client.chat.completions.create(
            model=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
            messages=[
                {"role": "system", "content": """
                    You are a creative prompt engineer for legal blog images. Create safe and professional image prompts that:
                    1. Are directly relevant to the blog content
                    2. Be 'unique to the blog's content', not generic or reusable for any legal article
                    3. Reflect the main topic, themes, or message of the blog post
                    4. Focus on modern, visually appealing representations
                    5. Must pass Azure content filters
                    6. Avoids sensitive content
                    The prompt should be detailed and specific, including:
                        - Main subject
                        - Style description
                        - Color palette
                        - Composition notes
                        - Mood/tone
                    - Is based on this blog content:
                """},
                {"role": "user", "content": text_prompt[:1000]}
            ],
            temperature=1
        )
        return response.choices[0].message.content

class FileManager:
    @staticmethod
    def list_articles():
        """
        List all DOCX files in the articles directory
        Returns:
            List of article filenames
        """
        articles = [f for f in os.listdir(Config.ARTICLES_DIR) if f.endswith('.docx')]
        return articles
    
    @staticmethod
    def get_article_metadata():
        """
        Read and parse the metadata.json file
        Returns:
            Dictionary of article metadata
        """
        metadata_path = os.path.join(Config.ARTICLES_DIR, 'metadata.json')
        try:
            with open(metadata_path, 'r', encoding='utf-8') as f:
                content = f.read()
                metadata = json.loads(content)
                # Convert list to dictionary for easier lookup
                result = {article['filename']: article for article in metadata['articles']}
                return result
        except (FileNotFoundError, json.JSONDecodeError, KeyError) as e:
            print(f"Error reading metadata: {str(e)}")
            return {}
    
    @staticmethod
    def read_docx(filename):
        """
        Read content from a DOCX file
        Args:
            filename: Name of the DOCX file (may be URL-encoded)
        Returns:
            Extracted text content
        """
        # URL-decode the filename first
        decoded_filename = unquote(filename)
        
        # For article filenames, we need to be more permissive
        # Check if the file exists in the articles directory
        filepath = os.path.join(Config.ARTICLES_DIR, decoded_filename)
        
        # Normalize the path to prevent path traversal
        normalized_path = os.path.normpath(filepath)
        
        # Simple path traversal check - look for .. in the path
        if '..' in normalized_path:
            raise ValueError("Path traversal detected")
        
        # Check if file exists
        if not os.path.exists(normalized_path):
            raise FileNotFoundError(f"Article file not found: {decoded_filename}")
        
        doc = Document(normalized_path)
        return "\n".join([para.text for para in doc.paragraphs])
    
    @staticmethod
    def save_content(content):
        """
        Save generated content to a file
        Args:
            content: Content to save
        Returns:
            Filename of the saved content
        """
        filename = f"blog_{int(time.time())}.txt"
        path = os.path.join(Config.GENERATED_DIR, filename)
        with open(path, 'w', encoding='utf-8') as f:
            f.write(content)
        return filename

    @staticmethod
    def generate_formatted_docx(content, title="Legal Blog"):
        """Generate DOCX with exact formatting from markdown"""
        doc = Document()

        # Custom styles (can be modified)
        styles = {
            'h1': {'font_size': 16, 'bold': True, 'color': RGBColor(0, 32, 96)},
            'h2': {'font_size': 14, 'bold': True, 'color': RGBColor(0, 64, 128)},
            'h3': {'font_size': 12, 'bold': True, 'italic': True},
            'bold': {'bold': True},
            'normal': {'font_size': 11}
        }
        
        def apply_style(run, style):
            """Helper function to apply formatting"""
            run.font.size = Pt(style.get('font_size', 11))
            run.font.bold = style.get('bold', False)
            run.font.italic = style.get('italic', False)
            if 'color' in style:
                run.font.color.rgb = style['color']
        
        # Process markdown content line by line
        lines = content.split('\n')
        for line in lines:

            if line.replace('-', '').strip() == '' and len(line) >= 3:
                continue

            # Detect formatting
            if line.startswith('# '):  # H1
                para = doc.add_heading(level=1)
                run = para.add_run(line[2:].strip())
                apply_style(run, styles['h1'])
                
            elif line.startswith('## '):  # H2
                para = doc.add_heading(level=2)
                run = para.add_run(line[3:].strip())
                apply_style(run, styles['h2'])
                
            elif line.startswith('### '):  # H3
                para = doc.add_heading(level=3)
                run = para.add_run(line[4:].strip())
                apply_style(run, styles['h3'])
                
            elif '**' in line:  # Bold text
                para = doc.add_paragraph()
                parts = re.split(r'(\*\*.+?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        apply_style(run, styles['bold'])
                    else:
                        para.add_run(part)
            
            else:  # Normal paragraph
                para = doc.add_paragraph()
                run = para.add_run(line)
                apply_style(run, styles['normal'])
        # Collect all empty paragraphs
        empty_paragraphs = [p for p in doc.paragraphs if not p.text.strip()]

        # Remove each empty paragraph from the document
        for p in empty_paragraphs:
            p._element.getparent().remove(p._element)

        # Save to bytes buffer
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

azure_services = AzureServices()
image_generator = ImageGenerator()

@app.template_filter('markdown')
def markdown_filter(text):
    html = markdown.markdown(text)
    soup = BeautifulSoup(html, 'html.parser')
    return str(soup)

@app.route('/')
def home():
    if not UserSession.get_current_user():
        return redirect(url_for('login'))
    return redirect(url_for('dashboard'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        firm = request.form['firm']
        location = request.form['location']
        lawyer_name = request.form['lawyer_name']
        state = request.form['state']
        address = request.form.get('address', '')
        planning_session = request.form.get('planning_session', '')
        other_planning_session = request.form.get('other_planning_session', '')
        discovery_call_link = request.form.get('discovery_call_link', '')

        if UserSession.register(email, password, firm, location, lawyer_name, state, 
                              address, planning_session, other_planning_session, discovery_call_link):
            # Auto-login after registration
            UserSession.login(email, password)
            return redirect(url_for('dashboard'))
        
        return render_template('register.html', error="Email already registered")
    
    return render_template('register.html')

@app.route('/profile', methods=['GET', 'POST'])
def profile():
    user = UserSession.get_current_user()
    if not user:
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        # Handle both form data and JSON requests
        if request.is_json:
            data = request.get_json()
            firm = data.get('firm', '')
            location = data.get('location', '')
            lawyer_name = data.get('lawyer_name', '')
            state = data.get('state', '')
            address = data.get('address', '')
            planning_session = data.get('planning_session', '')
            discovery_call_link = data.get('discovery_call_link', '')
            keywords = data.get('keywords', '')
        else:
            firm = request.form['firm']
            location = request.form['location']
            lawyer_name = request.form['lawyer_name']
            state = request.form['state']
            address = request.form.get('address', '')
            planning_session = request.form.get('planning_session', '')
            discovery_call_link = request.form.get('discovery_call_link', '')
            keywords = request.form.get('keywords', '')
        
        if UserSession.update_profile(user['username'], firm, location, lawyer_name, state, address, planning_session, "", discovery_call_link):
            session['user']['firm'] = firm
            session['user']['location'] = location
            session['user']['lawyer_name'] = lawyer_name
            session['user']['state'] = state
            session['user']['address'] = address
            session['user']['planning_session'] = planning_session
            session['user']['discovery_call_link'] = discovery_call_link
            session['user']['keywords'] = keywords
            session.modified = True
            
            # Return JSON response for AJAX requests
            if request.is_json:
                return jsonify({
                    'success': True,
                    'firm': firm,
                    'location': location,
                    'lawyer_name': lawyer_name,
                    'state': state,
                    'address': address,
                    'planning_session': planning_session,
                    'discovery_call_link': discovery_call_link
                })
            
            return redirect(url_for('dashboard'))
        
        # Return JSON response for AJAX requests
        if request.is_json:
            return jsonify({'success': False, 'error': 'Update failed'}), 400
        
        return render_template('profile.html', error="Update failed", user=session['user'])
    
    return render_template('profile.html', user=session['user'])

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        if UserSession.login(request.form['email'], request.form['password']):
            return redirect(url_for('dashboard'))
        return render_template('login.html', error="Invalid credentials")
    return render_template('login.html')

@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form['email']
        
        # Check if user exists
        db = get_db()
        user = db.execute('SELECT * FROM users WHERE email = ?', (email,)).fetchone()
        
        if user:
            # Generate reset token
            token = secrets.token_urlsafe(32)
            expires = datetime.now() + timedelta(hours=24)
            
            # Store reset token in database (SQL Server compatible)
            # First, delete any existing tokens for this email
            db.execute('DELETE FROM password_resets WHERE email = ?', (email,))
            
            # Then insert the new token
            db.execute('''
                INSERT INTO password_resets (email, token, expires)
                VALUES (?, ?, ?)
            ''', (email, token, expires))
            db.commit()
            
            # In a real application, you would send an email here
            # For now, we'll just show a success message with the token
            reset_url = url_for('reset_password', token=token, _external=True)
            
            return render_template('forgot_password.html', 
                                 success=f"Password reset link sent! For demo purposes, here's the link: {reset_url}")
        else:
            return render_template('forgot_password.html', 
                                 error="If an account with that email exists, a reset link has been sent.")
    
    return render_template('forgot_password.html')

@app.route('/reset_password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    db = get_db()
    
    # Check if token is valid and not expired
    reset_record = db.execute('''
        SELECT * FROM password_resets 
        WHERE token = ? AND expires > ? AND used = 0
    ''', (token, datetime.now())).fetchone()
    
    if not reset_record:
        return render_template('reset_password.html', token=token,
                             error="Invalid or expired reset link. Please request a new one.")
    
    if request.method == 'POST':
        password = request.form['password']
        confirm_password = request.form['confirm_password']
        
        if password != confirm_password:
            return render_template('reset_password.html', token=token,
                                 error="Passwords do not match.")
        
        if len(password) < 6:
            return render_template('reset_password.html', token=token,
                                 error="Password must be at least 6 characters long.")
        
        # Update user password - use index-based access for SQL Server cursor
        hashed_password = generate_password_hash(password)
        # reset_record[1] is the email field (index 1)
        db.execute('UPDATE users SET password = ? WHERE email = ?', 
                  (hashed_password, reset_record[1]))
        
        # Mark token as used
        db.execute('UPDATE password_resets SET used = 1 WHERE token = ?', (token,))
        db.commit()
        
        return render_template('reset_password.html', token=token,
                             success="Password reset successfully! You can now login with your new password.")
    
    return render_template('reset_password.html', token=token)

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/dashboard')
def dashboard():
    user = UserSession.get_current_user()
    if not user:
        return redirect(url_for('login'))
    
    # Get articles and their metadata
    articles = FileManager.list_articles()
    metadata = FileManager.get_article_metadata()
    
    # Get unique series names
    series_list = set()
    for article in articles:
        meta = metadata.get(article, {})
        if 'series' in meta:
            series_list.add(meta['series'])
    series_list = sorted(series_list) if series_list else None

    # Combine standard tones with user's custom tones
    standard_tones = [
        ('Professional', 'Formal and business-like tone suitable for corporate audiences'),
        ('Friendly', 'Warm and approachable tone that builds rapport with readers'),
        ('Educational', 'Clear and informative tone designed to explain concepts')
    ]
    
    custom_tones = user.get('custom_tones', [])
    all_tones = standard_tones + [(t['name'], t['description']) for t in custom_tones]
    
    # Convert to the format expected by the template
    tone_options = [t[0] for t in all_tones]
    tone_descriptions = {t[0]: t[1] for t in all_tones}
    
    return render_template('dashboard.html', 
                         user=user,
                         username=user['username'],
                         articles=articles,
                         metadata=metadata,
                         tone_options=tone_options,
                         tone_descriptions=tone_descriptions,
                         user_keywords=user.get('keywords', ''),
                         series_list=series_list)

@app.route('/add_tone', methods=['POST'])
def add_tone():
    user = UserSession.get_current_user()
    if not user:
        return jsonify({'success': False, 'error': 'Not logged in'}), 401
    
    data = request.get_json() if request.is_json else request.form
    tone_name = data.get('tone_name', '').strip()
    tone_description = data.get('tone_description', '').strip()
    
    if not tone_name:
        return jsonify({'success': False, 'error': 'Tone name is required'}), 400
    
    if UserSession.add_custom_tone(user['id'], tone_name, tone_description):
        return jsonify({
            'success': True,
            'tone_name': tone_name,
            'tone_description': tone_description
        })
    
    return jsonify({'success': False, 'error': 'Tone with this name already exists'}), 400

@app.route('/submit_feedback', methods=['POST'])
def submit_feedback():
    if 'user' not in session:
        return jsonify({'success': False, 'message': 'User not authenticated'})
    
    try:
        message = request.form.get('message')
        contact_email = request.form.get('contact_email')
        
        # Validate required fields
        if not message:
            return jsonify({'success': False, 'message': 'Please provide a message'})
        
        # Set default values for removed fields
        feedback_type = 'general'
        priority = 'medium'
        subject = 'User Feedback'
        
        # Submit feedback using UserSession method
        if UserSession.submit_feedback(session['user']['id'], feedback_type, priority, subject, message, contact_email):
            return jsonify({'success': True, 'message': 'Feedback submitted successfully'})
        else:
            return jsonify({'success': False, 'message': 'An error occurred while submitting feedback'})
        
    except Exception as e:
        print(f"Error submitting feedback: {str(e)}")
        return jsonify({'success': False, 'message': 'An error occurred while submitting feedback'})

@app.route('/select/<article>', methods=['GET', 'POST'])
def select_article(article):
    user = UserSession.get_current_user()
    if not user:
        return redirect(url_for('login'))
    
    # Initialize variables that will be used in both GET and POST
    firm = ''
    location = ''
    
    if request.method == 'POST':
        tone = request.form.get('tone')
        tone_description = request.form.get('toneDescription')
        custom_tone = request.form.get('customToneName')

        if tone == 'custom' and custom_tone:
            tone = custom_tone
            
        keywords = request.form.get('keywords', '')
        firm = request.form.get('firm', '')
        location = request.form.get('location', '')
        lawyer_name = user.get('lawyer_name', '')
        city = user.get('location', '')
        state = user.get('state', '')
        planning_session_name = request.form.get('planning_session_name','') 
        discovery_call_link = request.form.get('discovery_call_link','')
        if not planning_session_name:
            planning_session_name="Life & Legacy Planning Session"

        # Generate the blog post with the selected tone
        blog_content = azure_services.rewrite_content(
            FileManager.read_docx(article),
            tone,
            tone_description,
            keywords,
            firm,
            location,
            lawyer_name,
            city,
            state,
            planning_session_name,
            discovery_call_link
        )
        
        # Save the generated content to a file
        filename = FileManager.save_content(blog_content)
        
        # Set up the session data for the review page (without image initially)
        session['current_post'] = {
            'original': article,
            'content': blog_content,
            'image': None,  # Image will be generated later when requested
            'created': datetime.now().strftime("%Y-%m-%d %H:%M"),
            'tone': tone,
            'filename': filename
        }
        
        # Initialize chat history
        session['chat_history'] = [{
            'role': 'assistant',
            'content': blog_content,
            'content_is_blog': True,
            'timestamp': datetime.now().strftime("%H:%M:%S")
        }]
        
        # Generate a unique session ID for the chat
        session['session_id'] = os.urandom(16).hex()
        
        return redirect(url_for('review'))
    
    tone_options = [
        'Professional',
        'Friendly',
        'Educational'
    ]
    
    tone_descriptions = {
        'Professional': 'Formal and business-like tone suitable for corporate audiences',
        'Friendly': 'Warm and approachable tone that builds rapport with readers',
        'Educational': 'Clear and informative tone designed to explain concepts'
    }
    
    return render_template('select.html',
                         article_name=article,
                         tone_options=tone_options,
                         tone_descriptions=tone_descriptions,
                         firm=firm,
                         location=location)

@app.route('/use_version', methods=['POST'])
def use_version():
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    selected_content = request.form['content']
    
    session['current_post']['content'] = selected_content
    session.modified = True
    
    return redirect(url_for('finalize'))

@app.route('/finalize')
def finalize():
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    post = session['current_post']
    filename = FileManager.save_content(post['content'])
    image_url = url_for('static', filename=f'generated/{post["image"]}') if post.get('image') else None
    
    return render_template('finalize.html', 
                         post=post,
                         filename=filename,
                         image_url=image_url)

@app.route('/review', methods=['GET', 'POST'])
def review():
    # Check if we have a filename parameter but no current_post in session
    filename = request.args.get('filename')
    if filename and 'current_post' not in session:
        # Try to load the content from the file
        try:
            # Use safe path validation
            filepath = get_safe_file_path(Config.GENERATED_DIR, filename)
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
                
            # Set up the session data
            session['current_post'] = {
                'content': content,
                'filename': filename,
                'created': datetime.now().strftime("%Y-%m-%d %H:%M")
            }
            
            # Initialize chat history
            session['chat_history'] = [{
                'role': 'assistant',
                'content': content,
                'content_is_blog': True,
                'timestamp': datetime.now().strftime("%H:%M:%S")
            }]
            
            # Generate a unique session ID for the chat
            session['session_id'] = os.urandom(16).hex()
        except (ValueError, FileNotFoundError) as e:
            print(f"Error loading file: {e}")
            return redirect(url_for('dashboard'))
        except Exception as e:
            print(f"Error loading file: {e}")
            return redirect(url_for('dashboard'))
    
    # If we still don't have current_post in session, redirect to dashboard
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    post = session['current_post']
    
    if 'session_id' not in session:
        session['session_id'] = os.urandom(16).hex()
    
    if 'chat_history' not in session:
        session['chat_history'] = [{
            'role': 'assistant',
            'content': post['content'],
            'content_is_blog': True,
            'timestamp': datetime.now().strftime("%H:%M:%S")
        }]
    
    if request.method == 'POST':
        if 'edit_message' in request.form:  # Chat-style editing
            user_message = request.form['edit_message']
            
            current_content = next(
                (msg['content'] for msg in reversed(session['chat_history']) 
                 if msg['content_is_blog']),
                post['content']
            )
            
            edited_content = azure_services.edit_content(
                session['session_id'],
                user_message,
                current_content
            )
            
            session['chat_history'].append({
                'role': 'user',
                'content': user_message,
                'content_is_blog': False,
                'timestamp': datetime.now().strftime("%H:%M:%S")
            })
            session['chat_history'].append({
                'role': 'assistant',
                'content': edited_content,
                'content_is_blog': True,
                'timestamp': datetime.now().strftime("%H:%M:%S")
            })
            
            post['content'] = edited_content
            session['current_post'] = post
            
        elif 'content' in request.form:  # Manual editing
            post['content'] = request.form['content']
            session['current_post'] = post
            session['chat_history'].append({
                'role': 'assistant',
                'content': post['content'],
                'content_is_blog': True,
                'timestamp': datetime.now().strftime("%H:%M:%S")
            })

        session.modified = True
        return redirect(url_for('review'))
    
    # Save the current content to a file and get the filename
    if 'filename' not in post:
        filename = FileManager.save_content(post['content'])
        post['filename'] = filename
        session['current_post'] = post
    
    image_url = url_for('static', filename=f'generated/{post["image"]}') if post.get('image') else None
    
    return render_template('review.html', 
                         post=post,
                         chat_history=session['chat_history'],
                         image_url=image_url)

@app.route('/save_changes', methods=['POST'])
def save_changes():
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    edited_content = request.form.get('content', '')
    
    session['current_post']['content'] = edited_content
    
    if 'chat_history' not in session:
        session['chat_history'] = []
    
    session['chat_history'].append({
        'role': 'system',
        'content': 'User saved manual changes',
        'content_is_blog': False,
        'timestamp': datetime.now().strftime("%H:%M:%S")
    })
    
    session.modified = True
    return redirect(url_for('finalize'))

@app.route('/download/<filename>')
def download(filename):
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    try:
        # For generated files, use simple path validation since they're internally generated
        filepath = os.path.join(Config.GENERATED_DIR, filename)
        normalized_path = os.path.normpath(filepath)
        
        # Simple path traversal check - look for .. in the path
        if '..' in normalized_path:
            print(f"Path traversal detected: {normalized_path}")
            return redirect(url_for('review'))
        
        # Check if file exists
        if not os.path.exists(normalized_path):
            print(f"File not found: {normalized_path}")
            return redirect(url_for('review'))
        
        with open(normalized_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Get title
        title = session['current_post'].get('original', 'Legal Blog').replace('.docx', '')
        
        # Generate formatted DOCX
        docx_file = FileManager.generate_formatted_docx(content, title)
        
        return send_file(
            docx_file,
            as_attachment=True,
            download_name=f"{title}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"Download error: {e}")
        return redirect(url_for('review'))

@app.route('/generate_image')
def generate_image():
    if 'current_post' not in session:
        return redirect(url_for('dashboard'))
    
    # Generate image based on current content
    image_filename = image_generator.generate_image(session['current_post']['content'])
    
    if image_filename:
        session['current_post']['image'] = image_filename
        session.modified = True
    
    return redirect(url_for('review'))
    
@app.teardown_appcontext
def teardown_db(exception):
    close_db()

@app.route('/preview_article/<article>')
def preview_article(article):
    try:
        # URL-decode the article filename first
        decoded_article = unquote(article)
        
        # For article filenames, we need to be more permissive
        # Check if the file exists in the articles directory
        article_path = os.path.join(Config.ARTICLES_DIR, decoded_article)
        normalized_article_path = os.path.normpath(article_path)
        
        # Simple path traversal check - look for .. in the path
        if '..' in normalized_article_path:
            return jsonify({'error': 'Invalid article path'}), 400
        
        # Try to read markdown file first
        markdown_filename = decoded_article.replace('.docx', '.md')
        markdown_path = os.path.join(Config.ARTICLES_DIR, markdown_filename)
        normalized_markdown_path = os.path.normpath(markdown_path)
        
        # Ensure markdown path is also safe
        if '..' in normalized_markdown_path:
            return jsonify({'error': 'Invalid markdown path'}), 400
        
        if os.path.exists(normalized_markdown_path):
            # Read the markdown content
            with open(normalized_markdown_path, 'r', encoding='utf-8') as f:
                content = f.read()
        else:
            # If markdown doesn't exist, read from docx
            if not os.path.exists(normalized_article_path):
                return jsonify({'error': 'Article not found'}), 404
                
            doc = Document(normalized_article_path)
            content = "\n".join([para.text for para in doc.paragraphs])
            
        # Convert the content to HTML for preview
        html_content = markdown.markdown(content)
        return jsonify({'content': html_content})
    except (ValueError, FileNotFoundError) as e:
        print(f"File access error in preview_article: {str(e)}")
        return jsonify({'error': 'Article not found'}), 404
    except Exception as e:
        print(f"Error in preview_article: {str(e)}")  # Add logging
        return jsonify({'error': str(e)}), 500

def is_safe_filename(filename):
    """Validate that filename is safe and doesn't contain path traversal characters"""
    if not filename:
        return False
    
    # Check for path traversal characters
    dangerous_chars = ['..', '/', '\\', ':', '*', '?', '"', '<', '>', '|']
    for char in dangerous_chars:
        if char in filename:
            return False
    
    # Check if filename is only alphanumeric, dots, hyphens, and underscores
    if not re.match(r'^[a-zA-Z0-9._-]+$', filename):
        return False
    
    return True

def get_safe_file_path(base_dir, filename):
    """Safely construct a file path within the base directory"""
    if not is_safe_filename(filename):
        raise ValueError("Invalid filename")
    
    # Normalize the path to prevent path traversal
    full_path = os.path.normpath(os.path.join(base_dir, filename))
    
    # Ensure the path is within the base directory
    if not full_path.startswith(os.path.abspath(base_dir)):
        raise ValueError("Path traversal detected")
    
    return full_path

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8000, debug=True)